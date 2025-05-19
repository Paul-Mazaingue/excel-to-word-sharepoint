import os
import unittest
import tempfile
import pandas as pd
from pathlib import Path
from unittest.mock import patch, MagicMock

# Importer les fonctions à tester
from sharepoint_utils import (
    download_file_from_onedrive,
    extract_excel_row,
    check_file_exists_on_sharepoint,
    create_word_from_template,
    upload_file_to_sharepoint,
    create_temp_directory,
    clean_directory
)

class TestSharepointUtils(unittest.TestCase):
    """Tests pour les utilitaires SharePoint"""
    
    def setUp(self):
        """Configuration avant chaque test"""
        # Créer un répertoire temporaire pour les tests
        self.test_dir = Path(tempfile.mkdtemp(prefix="sharepoint_test_"))
        
        # Créer un fichier Excel de test
        self.test_excel_path = self.test_dir / "test_data.xlsx"
        df = pd.DataFrame({
            'nom': ['Test User'],
            'email': ['test@example.com'],
            'telephone': ['0123456789']
        })
        df.to_excel(self.test_excel_path, index=False)
        
        # Créer un template Word simple pour les tests
        self.test_word_template = self.test_dir / "test_template.docx"
        from docx import Document
        doc = Document()
        doc.add_paragraph("Nom: ${nom}")
        doc.add_paragraph("Email: ${email}")
        doc.save(self.test_word_template)
    
    def tearDown(self):
        """Nettoyage après chaque test"""
        clean_directory(self.test_dir)
    
    @patch('sharepoint_utils.subprocess.run')
    def test_download_file_from_onedrive(self, mock_run):
        """Test du téléchargement d'un fichier depuis OneDrive"""
        # Configurer le mock pour simuler un téléchargement réussi
        mock_process = MagicMock()
        mock_process.returncode = 0
        mock_process.stdout = "Téléchargement réussi"
        mock_run.return_value = mock_process
        
        # Créer un fichier fictif qui sera "téléchargé"
        test_file = self.test_dir / "test_download.txt"
        with open(test_file, "w") as f:
            f.write("contenu de test")
        
        # Simuler un téléchargement
        with patch('pathlib.Path.exists', return_value=True):
            with patch('sharepoint_utils.Path.name', return_value="test_download.txt"):
                result = download_file_from_onedrive("OneDrive", "dossier/test_download.txt", self.test_dir)
                
                # Vérifier que la fonction a appelé subprocess.run correctement
                mock_run.assert_called_once()
                self.assertIsNotNone(result)
    
    def test_extract_excel_row(self):
        """Test de l'extraction d'une ligne d'un fichier Excel"""
        # Tester avec un fichier Excel existant
        row = extract_excel_row(self.test_excel_path)
        
        # Vérifier que les données sont correctement extraites
        self.assertIsNotNone(row)
        self.assertEqual(row['nom'], 'Test User')
        self.assertEqual(row['email'], 'test@example.com')
        self.assertEqual(row['telephone'], '0123456789')
        
        # Tester avec un fichier inexistant
        nonexistent_file = self.test_dir / "nonexistent.xlsx"
        result = extract_excel_row(nonexistent_file)
        self.assertIsNone(result)
    
    @patch('sharepoint_utils.subprocess.run')
    def test_check_file_exists_on_sharepoint(self, mock_run):
        """Test de la vérification d'existence d'un fichier sur SharePoint"""
        # Cas 1: Le fichier existe
        mock_process_exists = MagicMock()
        mock_process_exists.returncode = 0
        mock_process_exists.stdout = "fichier.txt\n"
        mock_run.return_value = mock_process_exists
        
        exists_result = check_file_exists_on_sharepoint("SharePoint", "Documents", "fichier.txt")
        self.assertTrue(exists_result)
        
        # Cas 2: Le fichier n'existe pas
        mock_process_not_exists = MagicMock()
        mock_process_not_exists.returncode = 0
        mock_process_not_exists.stdout = ""
        mock_run.return_value = mock_process_not_exists
        
        not_exists_result = check_file_exists_on_sharepoint("SharePoint", "Documents", "fichier_inconnu.txt")
        self.assertFalse(not_exists_result)
    
    def test_create_word_from_template(self):
        """Test de création d'un document Word à partir d'un modèle"""
        # Créer des données de test
        data = pd.Series({
            'nom': 'Jean Dupont',
            'email': 'jean.dupont@example.com', 
            'telephone': '0123456789'
        })
        
        # Générer un document à partir du modèle
        output_path = self.test_dir / "output_doc.docx"
        result_path = create_word_from_template(self.test_word_template, data, output_path)
        
        # Vérifier que le document a été créé
        self.assertIsNotNone(result_path)
        self.assertTrue(output_path.exists())
        
        # Vérifier le contenu du document généré (vérification basique)
        from docx import Document
        doc = Document(result_path)
        text_content = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("Jean Dupont", text_content)
        self.assertIn("jean.dupont@example.com", text_content)
    
    @patch('sharepoint_utils.subprocess.run')
    @patch('sharepoint_utils.check_file_exists_on_sharepoint')
    def test_upload_file_to_sharepoint(self, mock_check, mock_run):
        """Test de téléversement d'un fichier vers SharePoint"""
        # Créer un fichier de test à téléverser
        test_file = self.test_dir / "upload_test.txt"
        with open(test_file, "w") as f:
            f.write("Contenu de test pour téléversement")
        
        # Configurer les mocks
        mock_check.return_value = False  # Simuler que le fichier n'existe pas déjà
        
        mock_process = MagicMock()
        mock_process.returncode = 0
        mock_run.return_value = mock_process
        
        # Tester le téléversement
        with patch('sharepoint_utils.check_file_exists_on_sharepoint', side_effect=[False, True]):
            result = upload_file_to_sharepoint(test_file, "SharePoint", "Documents")
            self.assertTrue(result)
            mock_run.assert_called()
    
    def test_create_temp_directory(self):
        """Test de création d'un répertoire temporaire"""
        temp_dir = create_temp_directory()
        self.assertIsNotNone(temp_dir)
        self.assertTrue(temp_dir.exists())
        self.assertTrue(temp_dir.is_dir())
        
        # Nettoyer
        clean_directory(temp_dir)
    
    def test_clean_directory(self):
        """Test de nettoyage d'un répertoire"""
        # Créer un répertoire avec des fichiers
        test_subdir = self.test_dir / "subdir"
        test_subdir.mkdir()
        
        test_file1 = test_subdir / "test1.txt"
        test_file2 = test_subdir / "test2.txt"
        
        with open(test_file1, "w") as f:
            f.write("test1")
        with open(test_file2, "w") as f:
            f.write("test2")
        
        # Vérifier que le répertoire et les fichiers existent
        self.assertTrue(test_subdir.exists())
        self.assertTrue(test_file1.exists())
        self.assertTrue(test_file2.exists())
        
        # Nettoyer et vérifier
        result = clean_directory(test_subdir)
        self.assertTrue(result)
        self.assertFalse(test_subdir.exists())

def run_tests():
    """Exécute les tests avec des options formatées"""
    import sys
    if len(sys.argv) > 1:
        # Si des arguments sont fournis, exécuter seulement les tests spécifiés
        test_names = sys.argv[1:]
        suite = unittest.TestSuite()
        for test_name in test_names:
            suite.addTest(TestSharepointUtils(test_name))
        unittest.TextTestRunner(verbosity=2).run(suite)
    else:
        # Sinon, exécuter tous les tests
        unittest.main(argv=['first-arg-is-ignored'], verbosity=2)

if __name__ == "__main__":
    print("=== Tests des utilitaires SharePoint ===")
    run_tests()
