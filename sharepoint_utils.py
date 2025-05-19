import os
import subprocess
from pathlib import Path
import shutil
import time
import tempfile
import uuid
import pandas as pd
from docx import Document
from datetime import datetime
import logging
import platform
import shutil

# Configuration de base du logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("sharepoint_utils")

# Détecte l'OS
if platform.system() == "Windows":
    RCLONE_PATH = "C:\\rclone\\rclone.exe"
else:
    RCLONE_PATH = shutil.which("rclone")

print("---------------------------------------------")
print("Système d'exploitation détecté:", platform.system())
print("Chemin de rclone:", RCLONE_PATH)
print("---------------------------------------------")

if not RCLONE_PATH or not Path(RCLONE_PATH).exists():
    raise FileNotFoundError(f"rclone introuvable à l'emplacement attendu ou dans le PATH : {RCLONE_PATH}")

def create_temp_directory():
    """Crée un répertoire temporaire unique pour les opérations de fichiers"""
    unique_id = uuid.uuid4().hex[:8]
    temp_base_dir = Path(tempfile.gettempdir())
    temp_dir = temp_base_dir / f"sharepoint_op_{unique_id}"
    temp_dir.mkdir(exist_ok=True)
    logger.info(f"Répertoire temporaire créé: {temp_dir}")
    return temp_dir

def clean_directory(directory):
    """Nettoie un répertoire avec plusieurs tentatives
    
    Args:
        directory (Path): Le répertoire à nettoyer
        
    Returns:
        bool: True si le nettoyage a réussi, False sinon
    """
    max_attempts = 3
    for attempt in range(max_attempts):
        try:
            # Attendre un moment pour que les processus libèrent les fichiers
            time.sleep(1)
            
            if not directory.exists():
                return True
                
            # Tenter de supprimer les fichiers individuellement d'abord
            for item in directory.glob('**/*'):
                if item.is_file():
                    try:
                        os.chmod(item, 0o666)  # S'assurer que le fichier est accessible
                        item.unlink()
                    except Exception as e:
                        logger.error(f"Impossible de supprimer {item}: {e}")
            
            # Puis essayer de supprimer les dossiers
            shutil.rmtree(directory, ignore_errors=True)
            
            if not directory.exists():
                return True
                
        except Exception as e:
            logger.warning(f"Tentative {attempt+1} de nettoyage échouée: {e}")
            
    return False

def download_file_from_onedrive(remote_name, remote_file_path, local_dir=None):
    """Télécharge un fichier depuis OneDrive vers un répertoire local
    
    Args:
        remote_name (str): Nom du remote rclone configuré (ex: "Onedrive")
        remote_file_path (str): Chemin du fichier dans OneDrive
        local_dir (Path, optional): Répertoire local où stocker le fichier.
                                   Si None, un répertoire temporaire sera créé.
    
    Returns:
        Path: Le chemin local du fichier téléchargé ou None en cas d'échec
    """
    try:
        # Créer un répertoire temporaire si nécessaire
        if local_dir is None:
            local_dir = create_temp_directory()
        elif isinstance(local_dir, str):
            local_dir = Path(local_dir)
            
        local_dir.mkdir(exist_ok=True, parents=True)
        
        # Préparer les arguments pour rclone
        args = [
            RCLONE_PATH,
            "copy",
            f"{remote_name}:{remote_file_path}",
            str(local_dir),
            "--drive-chunk-size", "64M"
        ]
        
        logger.info(f"Téléchargement du fichier {remote_file_path} depuis {remote_name}...")
        logger.debug(f"Commande: {' '.join(args)}")
        
        result = subprocess.run(args, capture_output=True, text=True)
        
        if result.returncode != 0:
            logger.error(f"Erreur lors du téléchargement: {result.stderr}")
            return None
            
        # Récupérer le nom du fichier à partir du chemin distant
        filename = Path(remote_file_path).name
        local_file_path = local_dir / filename
        
        if local_file_path.exists():
            logger.info(f"Fichier téléchargé avec succès: {local_file_path}")
            return local_file_path
        else:
            logger.error(f"Le fichier téléchargé n'a pas été trouvé à l'emplacement attendu: {local_file_path}")
            return None
            
    except Exception as e:
        logger.error(f"Erreur lors du téléchargement depuis OneDrive: {str(e)}")
        return None

def extract_excel_row(excel_file_path, row_index=-1):
    """Extrait une ligne spécifique d'un fichier Excel
    
    Args:
        excel_file_path (Path): Chemin vers le fichier Excel
        row_index (int): Index de la ligne à extraire (-1 pour la dernière ligne)
    
    Returns:
        pd.Series: La ligne extraite sous forme de Series pandas ou None en cas d'échec
    """
    try:
        # Vérifier que le fichier existe
        if not Path(excel_file_path).exists():
            logger.error(f"Le fichier Excel n'existe pas: {excel_file_path}")
            return None
            
        # Charger le fichier Excel en forçant les colonnes comme texte
        df = pd.read_excel(excel_file_path, dtype=str)
        
        # Vérifier que le DataFrame n'est pas vide
        if df.empty:
            logger.error(f"Le fichier Excel {excel_file_path} est vide")
            return None
            
        # Vérifier que l'index demandé est valide
        if abs(row_index) > len(df):
            logger.error(f"Index de ligne invalide: {row_index}, max: {len(df)}")
            return None
            
        # Extraire la ligne (déjà en str grâce à dtype=str) et remplacer NaN par ''
        row_data = df.iloc[row_index].fillna('')

        logger.info(f"Ligne extraite du fichier Excel avec succès: {row_index}")
        return row_data
        
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction de données Excel: {str(e)}")
        return None


def check_file_exists_on_sharepoint(remote_name, remote_folder, filename):
    """Vérifie si un fichier existe sur SharePoint
    
    Args:
        remote_name (str): Nom du remote rclone configuré
        remote_folder (str): Dossier dans le remote SharePoint
        filename (str): Nom du fichier à vérifier
    
    Returns:
        bool: True si le fichier existe, False sinon
    """
    try:
        # Construire le chemin complet
        remote_path = f"{remote_name}:{remote_folder}/{filename}"
        
        # Préparer la commande pour vérifier l'existence du fichier
        check_args = [RCLONE_PATH, "lsf", remote_path]
        logger.debug(f"Commande de vérification: {' '.join(check_args)}")
        
        # Exécuter la commande
        check_result = subprocess.run(check_args, capture_output=True, text=True)
        
        # Analyser le résultat
        if check_result.returncode == 0 and check_result.stdout.strip():
            logger.info(f"Le fichier {filename} existe sur SharePoint")
            return True
        else:
            logger.info(f"Le fichier {filename} n'existe pas sur SharePoint")
            return False
            
    except Exception as e:
        logger.error(f"Erreur lors de la vérification du fichier sur SharePoint: {str(e)}")
        return False

def delete_file_from_sharepoint(remote_name, remote_folder, filename):
    """Supprime un fichier de SharePoint
    
    Args:
        remote_name (str): Nom du remote rclone configuré
        remote_folder (str): Dossier dans le remote SharePoint
        filename (str): Nom du fichier à supprimer
    
    Returns:
        bool: True si la suppression a réussi, False sinon
    """
    try:
        # Construire le chemin complet
        remote_path = f"{remote_name}:{remote_folder}/{filename}"
        
        # Vérifier d'abord que le fichier existe
        if not check_file_exists_on_sharepoint(remote_name, remote_folder, filename):
            logger.info(f"Pas de suppression nécessaire - le fichier {filename} n'existe pas sur SharePoint")
            return True
            
        # Préparer la commande de suppression
        delete_args = [RCLONE_PATH, "deletefile", remote_path]
        logger.debug(f"Commande de suppression: {' '.join(delete_args)}")
        
        # Exécuter la commande
        delete_result = subprocess.run(delete_args, capture_output=True, text=True)
        
        # Analyser le résultat
        if delete_result.returncode == 0:
            logger.info(f"Fichier {filename} supprimé avec succès de SharePoint")
            # Attendre un peu pour s'assurer que la suppression est prise en compte
            time.sleep(2)
            return True
        else:
            logger.error(f"Impossible de supprimer le fichier {filename} de SharePoint: {delete_result.stderr}")
            return False
            
    except Exception as e:
        logger.error(f"Erreur lors de la suppression du fichier sur SharePoint: {str(e)}")
        return False

def create_word_from_template(template_path, data_row, output_path=None):
    """Crée un document Word à partir d'un modèle et de données
    
    Args:
        template_path (Path): Chemin vers le modèle Word
        data_row (pd.Series): Données pour remplir le modèle
        output_path (Path, optional): Chemin pour le fichier de sortie.
                                     Si None, un fichier est créé dans un répertoire temporaire.
    
    Returns:
        Path: Le chemin du document Word généré ou None en cas d'échec
    """
    try:
        # Vérifier que le modèle existe
        template_path = Path(template_path)
        if not template_path.exists():
            logger.error(f"Le modèle Word n'existe pas: {template_path}")
            return None
            
        # Créer le chemin de sortie si nécessaire
        if output_path is None:
            temp_dir = create_temp_directory()
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = temp_dir / f"document_{timestamp}.docx"
        elif isinstance(output_path, str):
            output_path = Path(output_path)
            
        # Créer le dossier parent du fichier de sortie si nécessaire
        output_path.parent.mkdir(exist_ok=True, parents=True)
            
        # Charger le modèle Word
        doc = Document(template_path)
        
        # Définir l'espace de noms XML pour Word
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # Fonction de normalisation des chaînes
        def normalize_string(text):
            if not isinstance(text, str):
                return str(text).lower()
            # Supprimer les accents
            import unicodedata
            text = unicodedata.normalize('NFKD', text).encode('ASCII', 'ignore').decode('ASCII')
            # Convertir en minuscules
            text = text.lower()
            # Remplacer la ponctuation par des espaces
            import string
            for char in string.punctuation:
                text = text.replace(char, ' ')
            # Réduire les espaces multiples
            text = ' '.join(text.split())
            return text
        
        # Fonction pour formater les dates
        def format_date(value):
            if not isinstance(value, str):
                value = str(value)
            
            # Détecter le format de date YYYY-MM-DD HH:MM:SS
            import re
            date_pattern = r'^\d{4}-\d{2}-\d{2}\s\d{2}:\d{2}:\d{2}$'
            if re.match(date_pattern, value):
                try:
                    # Convertir la date au format jj-mm-yyyy
                    date_obj = datetime.strptime(value, '%Y-%m-%d %H:%M:%S')
                    return date_obj.strftime('%d-%m-%Y')
                except Exception as e:
                    logger.warning(f"Erreur de conversion de date '{value}': {e}")
                    return value
            return value
        
        # Créer un dictionnaire de correspondance entre clés normalisées et clés originales
        normalized_keys = {}
        for key in data_row.index:
            normalized_keys[normalize_string(key)] = key
        
        logger.debug(f"Clés normalisées: {normalized_keys}")
        
        # Méthode 1: Remplacement des contrôles de contenu (Content Controls)
        try:
            # Dictionnaire pour suivre le nombre d'occurrences de chaque tag
            tag_counters = {tag: 0 for tag in data_row.index}
            
            # Parcourir toutes les parties du document
            for part in [doc.part]:
                # Obtenir l'élément XML racine
                root = part.element
                
                # Chercher tous les contrôles de contenu
                for sdt in root.findall('.//w:sdt', ns):
                    # Chercher le tag (étiquette) du contrôle
                    tag_elem = sdt.find('.//w:tag', ns)
                    if tag_elem is not None and '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val' in tag_elem.attrib:
                        tag = tag_elem.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val']
                        
                        # Normaliser le tag pour la comparaison
                        normalized_tag = normalize_string(tag)
                        logger.debug(f"Tag original: '{tag}', normalisé: '{normalized_tag}'")
                        
                        # Vérifier si le tag normalisé correspond à une clé normalisée
                        if normalized_tag in normalized_keys:
                            # Obtenir la clé originale
                            original_key = normalized_keys[normalized_tag]
                            # Obtenir la valeur depuis les données et la formater si c'est une date
                            value = str(format_date(data_row[original_key]))
                            
                            logger.debug(f"Correspondance trouvée - Tag: '{tag}' -> Clé: '{original_key}', Valeur: '{value}'")
                            
                            # Trouver l'élément de texte dans le contrôle et le remplacer
                            for text_elem in sdt.findall('.//w:t', ns):
                                text_elem.text = value
                                tag_counters[original_key] += 1
                                break  # Ne remplacer que le premier élément de texte dans ce contrôle
        except Exception as e:
            logger.warning(f"Erreur lors du traitement des contrôles de contenu: {str(e)}", exc_info=True)
        
        # Méthode 2: Remplacement des marqueurs de texte
        try:
            # Réinitialiser les compteurs pour les marqueurs de texte
            tag_counters = {tag: 0 for tag in data_row.index}
            
            # Fonction pour remplacer les tags dans un run
            def process_run(run):
                for key in data_row.index:
                    placeholder = f"${{{key}}}"
                    while placeholder in run.text:
                        # Remplacer seulement la première occurrence
                        tag_counters[key] += 1
                        # Formater la valeur si c'est une date
                        value = format_date(data_row[key])
                        parts = run.text.split(placeholder, 1)  # Split seulement sur la première occurrence
                        run.text = parts[0] + str(value) + parts[1] if len(parts) > 1 else parts[0]
            
            # Parcourir les paragraphes
            for p in doc.paragraphs:
                for run in p.runs:
                    process_run(run)
            
            # Également vérifier les tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                process_run(run)
        except Exception as e:
            logger.warning(f"Erreur lors du remplacement de texte: {str(e)}")
        
        # Sauvegarder le document Word rempli
        doc.save(output_path)
        logger.info(f"Document Word généré avec succès: {output_path}")
        return output_path
            
    except Exception as e:
        logger.error(f"Erreur lors de la création du document Word: {str(e)}")
        return None

def upload_file_to_sharepoint(local_file_path, remote_name, remote_folder, overwrite=True):
    """Téléverse un fichier sur SharePoint
    
    Args:
        local_file_path (Path): Chemin local du fichier à téléverser
        remote_name (str): Nom du remote rclone configuré
        remote_folder (str): Dossier dans le remote SharePoint
        overwrite (bool): Si True, écrase le fichier s'il existe déjà
    
    Returns:
        bool: True si le téléversement a réussi, False sinon
    """
    try:
        # Vérifier que le fichier local existe
        local_file_path = Path(local_file_path)
        if not local_file_path.exists():
            logger.error(f"Le fichier local n'existe pas: {local_file_path}")
            return False
            
        # Obtenir le nom du fichier
        filename = local_file_path.name
        
        # Construire le chemin distant
        remote_path = f"{remote_name}:{remote_folder}"
        remote_file = f"{remote_path}/{filename}"
        
        # Vérifier si le fichier existe déjà et le supprimer si nécessaire
        if overwrite and check_file_exists_on_sharepoint(remote_name, remote_folder, filename):
            if not delete_file_from_sharepoint(remote_name, remote_folder, filename):
                logger.warning(f"Impossible de supprimer le fichier existant. Le téléversement pourrait échouer.")
        
        # Créer un dossier temporaire pour la copie du fichier
        temp_dir = create_temp_directory()
        copy_file = temp_dir / filename
        
        try:
            # Copier le fichier dans le dossier temporaire
            shutil.copy2(local_file_path, copy_file)
            logger.debug(f"Fichier copié vers {copy_file}")
        except Exception as copy_error:
            logger.error(f"Erreur lors de la copie du fichier: {copy_error}")
            return False
        
        # Téléverser le fichier
        logger.info(f"Téléversement du fichier {filename} vers SharePoint...")
        args = [
            RCLONE_PATH,
            "copy",
            str(copy_file),
            remote_path + "/",
            "--ignore-checksum",
            "--ignore-size"
        ]
        
        logger.debug(f"Commande: {' '.join(args)}")
        result = subprocess.run(args, capture_output=True, text=True)
        
        success = False
        if result.returncode == 0:
            logger.info("Téléversement réussi!")
            success = True
        else:
            logger.warning(f"Première tentative de téléversement échouée: {result.stderr}")
            
            # Si l'upload échoue, tenter avec copyto pour cibler exactement le fichier
            logger.info("Tentative avec copyto...")
            copyto_args = [
                RCLONE_PATH,
                "copyto",
                str(copy_file),
                remote_file,
                "--ignore-checksum",
                "--ignore-size"
            ]
            
            result_retry = subprocess.run(copyto_args, capture_output=True, text=True)
            
            if result_retry.returncode == 0:
                logger.info("Téléversement réussi avec copyto!")
                success = True
            else:
                logger.error(f"Toutes les tentatives de téléversement ont échoué: {result_retry.stderr}")
                success = False
        
        # Vérifier si le fichier existe sur le serveur distant
        if success:
            time.sleep(1)  # Attendre un peu pour s'assurer que l'upload est terminé
            if check_file_exists_on_sharepoint(remote_name, remote_folder, filename):
                logger.info(f"✅ Vérification réussie: Le fichier {filename} existe sur SharePoint!")
            else:
                logger.warning(f"❌ Vérification échouée: Le fichier {filename} n'est pas visible sur SharePoint.")
                success = False
        
        # Nettoyer le répertoire temporaire
        clean_directory(temp_dir)
        
        return success
            
    except Exception as e:
        logger.error(f"Erreur lors du téléversement sur SharePoint: {str(e)}")
        return False
